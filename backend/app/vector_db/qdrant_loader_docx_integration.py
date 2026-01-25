# qdrant_docx_module.py
import uuid
import os
import re
import time
import hashlib
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional

import docx
from qdrant_client import QdrantClient, models
from qdrant_client.http.models import PointStruct, VectorParams, Distance
import numpy as np
from sentence_transformers import SentenceTransformer
import tqdm

# Конфигурация
QDRANT_HOST = "localhost"
QDRANT_PORT = 6333
COLLECTION_NAME = "ida_edubot"
VECTOR_SIZE = 1024
EMBEDDING_MODEL_NAME = "intfloat/multilingual-e5-large"

@dataclass
class DocumentMetadata:
    """Метаданные документа"""
    course_title: str = ""
    discipline_name: str = ""
    discipline_id: str = ""
    material_id: str = ""
    course_id: str = ""
    content_type: str = "document"
    difficulty: str = "medium"

@dataclass
class UploadResult:
    """Результат загрузки документа"""
    success: bool
    points_uploaded: int
    total_points: int
    collection_name: str
    metadata: DocumentMetadata
    message: str = ""

class QdrantDocxUploader:
    """Модуль для загрузки документов DOCX в Qdrant"""
    
    def __init__(self, qdrant_host: str = QDRANT_HOST, qdrant_port: int = QDRANT_PORT, 
                 collection_name: str = COLLECTION_NAME, model_name: str = EMBEDDING_MODEL_NAME):
        """Инициализация загрузчика"""
        self.qdrant_host = qdrant_host
        self.qdrant_port = qdrant_port
        self.collection_name = collection_name
        
        print("🧠 Загрузка модели эмбеддингов...")
        try:
            self.embedding_model = SentenceTransformer(model_name)
            print(f"✅ Модель загружена: {model_name}")
            
            # Тестируем размерность
            test_embedding = self.embedding_model.encode(["тест"])[0]
            self.model_dim = len(test_embedding)
            print(f"📏 Размерность модели: {self.model_dim}")
            
            if self.model_dim != VECTOR_SIZE:
                print(f"⚠️  Размерность модели ({self.model_dim}) не совпадает с конфигурацией ({VECTOR_SIZE})")
        except Exception as e:
            print(f"❌ Ошибка загрузки модели: {e}")
            raise
        
        # Подключение к Qdrant
        print(f"🔗 Подключение к Qdrant {qdrant_host}:{qdrant_port}...")
        try:
            self.client = QdrantClient(host=qdrant_host, port=qdrant_port, timeout=10)
            print(f"✅ Подключение к Qdrant установлено")
            
        except Exception as e:
            print(f"❌ Ошибка подключения к Qdrant: {e}")
            raise
    
    def _pad_or_truncate_vector(self, vector: np.ndarray) -> np.ndarray:
        """Дополнение или усечение вектора до VECTOR_SIZE"""
        if len(vector) > VECTOR_SIZE:
            return vector[:VECTOR_SIZE]
        elif len(vector) < VECTOR_SIZE:
            padded = np.zeros(VECTOR_SIZE)
            padded[:len(vector)] = vector
            return padded
        else:
            return vector
    
    def _get_embeddings(self, texts: List[str]) -> List[np.ndarray]:
        """Получение эмбеддингов с приведением к нужной размерности"""
        embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
        processed_embeddings = [self._pad_or_truncate_vector(emb) for emb in embeddings]
        return processed_embeddings
    
    def _clean_text_preserve_structure(self, text: str) -> str:
        """Очистка текста с сохранением структуры"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?;:()\-—\"\'\nа-яА-ЯёЁ]', '', text)
        text = re.sub(r'\s*([.,!?;:])\s*', r'\1 ', text)
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Разбивает текст на предложения"""
        sentences = []
        current_sentence = []
        i = 0
        
        while i < len(text):
            char = text[i]
            current_sentence.append(char)
            
            if char in '.!?':
                if char == '.':
                    abbreviations = ['т.д.', 'т.п.', 'др.', 'проф.', 'акад.', 'ст.', 'гл.',
                                    'рис.', 'табл.', 'см.', 'напр.', 'д.', 'и т.д.', 'и т.п.']
                    
                    is_abbreviation = False
                    for abbr in abbreviations:
                        if text[i-len(abbr)+1:i+1] == abbr:
                            is_abbreviation = True
                            break
                    
                    if not is_abbreviation and i + 1 < len(text) and text[i+1].isspace():
                        sentence = ''.join(current_sentence).strip()
                        if sentence:
                            sentences.append(sentence)
                        current_sentence = []
                else:
                    if i + 1 < len(text) and text[i+1].isspace():
                        sentence = ''.join(current_sentence).strip()
                        if sentence:
                            sentences.append(sentence)
                        current_sentence = []
            
            i += 1
        
        if current_sentence:
            sentence = ''.join(current_sentence).strip()
            if sentence:
                sentences.append(sentence)
        
        return sentences
    
    def _chunk_text_intelligent(self, text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
        """
        Умное разбиение текста на чанки с сохранением целостности предложений.
        """
        if not text or len(text.strip()) == 0:
            return []
        
        text = self._clean_text_preserve_structure(text)
        
        if len(text) <= chunk_size:
            return [text]
        
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            starts_with_lower = sentence and sentence[0].islower()
            
            if not current_chunk and starts_with_lower:
                current_chunk = sentence
                continue
            
            if len(current_chunk) + len(sentence) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Объединяем слишком маленькие чанки
        merged_chunks = []
        for chunk in chunks:
            if len(chunk) < 100 and merged_chunks:
                merged_chunks[-1] = merged_chunks[-1] + " " + chunk
            else:
                merged_chunks.append(chunk)
        
        # Добавляем перекрытие между чанками
        if len(merged_chunks) > 1:
            final_chunks = [merged_chunks[0]]
            
            for i in range(1, len(merged_chunks)):
                prev_chunk = merged_chunks[i-1]
                current_chunk = merged_chunks[i]
                
                if len(prev_chunk) > overlap:
                    overlap_text = prev_chunk[-overlap:]
                    for delimiter in ['. ', '! ', '? ', '; ']:
                        pos = overlap_text.find(delimiter)
                        if pos != -1:
                            overlap_text = overlap_text[pos+2:]
                            break
                else:
                    overlap_text = prev_chunk
                
                new_chunk = overlap_text + " " + current_chunk
                final_chunks.append(new_chunk.strip())
            
            return final_chunks
        
        return merged_chunks
    
    def _read_docx(self, file_path: str) -> str:
        """Чтение docx файла"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"❌ Ошибка чтения docx: {e}")
            raise
    
    def _get_simple_id(self, text: str) -> str:
        """Простая генерация ID из строки"""
        hash_obj = hashlib.md5(text.lower().strip().encode())
        hex_dig = hash_obj.hexdigest()
        return f"{hex_dig[:8]}-{hex_dig[8:12]}-{hex_dig[12:16]}-{hex_dig[16:20]}-{hex_dig[20:]}"
    
    def ensure_collection_exists(self) -> bool:
        """Проверяет существование коллекции и создает ее при необходимости"""
        try:
            collections = self.client.get_collections().collections
            collection_names = [c.name for c in collections]
            
            if self.collection_name in collection_names:
                print(f"✅ Коллекция '{self.collection_name}' существует")
                return True
            else:
                print(f"📦 Коллекция '{self.collection_name}' не найдена, создаю...")
                return self.create_collection()
                
        except Exception as e:
            print(f"❌ Ошибка проверки коллекции: {e}")
            return False
    
    def create_collection(self) -> bool:
        """Создание новой коллекции в Qdrant"""
        try:
            print(f"\n🔧 СОЗДАНИЕ КОЛЛЕКЦИИ '{self.collection_name}'")
            print(f"   Размерность: {self.model_dim}")
            print(f"   Метрика: {Distance.COSINE}")
            
            vectors_config = VectorParams(
                size=self.model_dim,
                distance=Distance.COSINE
            )
            
            print(f"🚀 Создание коллекции...")
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=vectors_config,
            )
            
            # Создаем индекс для payload
            self._create_payload_indexes()
            
            print(f"✅ Коллекция '{self.collection_name}' успешно создана!")
            return True
                
        except Exception as e:
            print(f"❌ Ошибка создания коллекции: {e}")
            return False
    
    def _create_payload_indexes(self):
        """Создание индексов для полей payload для ускорения поиска"""
        try:
            print("🔍 Создание индексов для ускорения поиска...")
            
            index_fields = [
                ("discipline_id", models.PayloadSchemaType.KEYWORD),
                ("course_id", models.PayloadSchemaType.KEYWORD),
                ("material_id", models.PayloadSchemaType.KEYWORD),
                ("discipline_name", models.PayloadSchemaType.TEXT),
                ("course_title", models.PayloadSchemaType.TEXT),
                ("difficulty", models.PayloadSchemaType.KEYWORD),
                ("content_type", models.PayloadSchemaType.KEYWORD),
                ("source_file", models.PayloadSchemaType.TEXT),
            ]
            
            for field_name, schema_type in index_fields:
                try:
                    self.client.create_payload_index(
                        collection_name=self.collection_name,
                        field_name=field_name,
                        field_schema=schema_type,
                    )
                    print(f"   ✓ Индекс для поля '{field_name}' создан")
                except Exception as e:
                    if "already exists" in str(e):
                        print(f"   ℹ️  Индекс для '{field_name}' уже существует")
                    else:
                        print(f"   ⚠️  Не удалось создать индекс для '{field_name}': {str(e)[:100]}")
            
            print("✅ Индексы созданы")
            
        except Exception as e:
            print(f"⚠️  Ошибка при создании индексов: {e}")
    
    def upload_document(self, doc_path: str, metadata: Optional[DocumentMetadata] = None) -> UploadResult:
        """Загрузка документа в коллекцию"""
        if not os.path.exists(doc_path):
            return UploadResult(
                success=False,
                points_uploaded=0,
                total_points=0,
                collection_name=self.collection_name,
                metadata=metadata or DocumentMetadata(),
                message=f"Файл не найден: {doc_path}"
            )
        
        print(f"\n📄 Обработка документа: {Path(doc_path).name}")
        
        # Используем переданные метаданные или создаем на основе имени файла
        if metadata is None:
            metadata = self._create_metadata_from_filename(doc_path)
        
        # Проверяем/создаем коллекцию
        if not self.ensure_collection_exists():
            return UploadResult(
                success=False,
                points_uploaded=0,
                total_points=0,
                collection_name=self.collection_name,
                metadata=metadata,
                message="Не удалось создать или найти коллекцию"
            )
        
        # Чтение документа
        print("📖 Чтение документа...")
        try:
            full_text = self._read_docx(doc_path)
            print(f"   Прочитано символов: {len(full_text):,}")
            
            if len(full_text.strip()) < 50:
                return UploadResult(
                    success=False,
                    points_uploaded=0,
                    total_points=0,
                    collection_name=self.collection_name,
                    metadata=metadata,
                    message="Текст слишком короткий"
                )
        except Exception as e:
            return UploadResult(
                success=False,
                points_uploaded=0,
                total_points=0,
                collection_name=self.collection_name,
                metadata=metadata,
                message=f"Ошибка чтения файла: {e}"
            )
        
        # Разбивка на чанки
        print("✂️  Интеллектуальное разбиение на чанки...")
        chunks = self._chunk_text_intelligent(full_text, chunk_size=800, overlap=200)
        print(f"   Создано чанков: {len(chunks)}")
        
        if not chunks:
            return UploadResult(
                success=False,
                points_uploaded=0,
                total_points=0,
                collection_name=self.collection_name,
                metadata=metadata,
                message="Не удалось создать чанки"
            )
        
        # Генерация эмбеддингов
        print("🔢 Генерация эмбеддингов...")
        try:
            embeddings = self._get_embeddings(chunks)
            print(f"   Сгенерировано эмбеддингов: {len(embeddings)}")
        except Exception as e:
            return UploadResult(
                success=False,
                points_uploaded=0,
                total_points=0,
                collection_name=self.collection_name,
                metadata=metadata,
                message=f"Ошибка генерации эмбеддингов: {e}"
            )
        
        # Подготовка точек
        print("📦 Подготовка данных...")
        points = []
        
        try:
            existing_count = self.client.count(
                collection_name=self.collection_name,
                exact=True
            ).count
            start_id = existing_count + 1
        except:
            start_id = 1
        
        for idx, (chunk_text, embedding) in enumerate(tqdm.tqdm(zip(chunks, embeddings), 
                                                                total=len(chunks), 
                                                                desc="Подготовка")):
            point_id = start_id + idx
            
            chunk_hash = hashlib.md5(f"{metadata.material_id}_{idx}".encode()).hexdigest()[:8]
            chunk_id_int = int(chunk_hash, 16) % 1000000
            
            point = PointStruct(
                id=point_id,
                vector=embedding.tolist(),
                payload={
                    "content_type": metadata.content_type,
                    "discipline_id": metadata.discipline_id,
                    "chunk_id": chunk_id_int,
                    "chunk_index": idx,
                    "course_title": metadata.course_title,
                    "difficulty": metadata.difficulty,
                    "material_id": metadata.material_id,
                    "course_id": metadata.course_id,
                    "discipline_name": metadata.discipline_name,
                    "chunk_text": chunk_text[:5000],
                    "source_file": os.path.basename(doc_path),
                    "total_chunks": len(chunks),
                    "upload_timestamp": int(time.time())
                }
            )
            points.append(point)
        
        # Загрузка в Qdrant
        print(f"\n🚀 Загрузка {len(points)} точек в Qdrant...")
        
        batch_size = 10
        success_count = 0
        
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            try:
                self.client.upsert(
                    collection_name=self.collection_name,
                    points=batch,
                    wait=True
                )
                success_count += len(batch)
                print(f"   ✓ Батч {i//batch_size + 1}: загружено {min(i + batch_size, len(points))}/{len(points)}")
            except Exception as e:
                print(f"   ✗ Ошибка батча {i//batch_size + 1}: {str(e)[:200]}")
                # Пробуем загрузить по одной
                for j, point in enumerate(batch):
                    try:
                        self.client.upsert(
                            collection_name=self.collection_name,
                            points=[point],
                            wait=False
                        )
                        success_count += 1
                    except Exception as point_error:
                        print(f"      ✗ Ошибка точки {i + j}: {str(point_error)[:100]}")
        
        # Итоги
        print("\n" + "="*50)
        if success_count > 0:
            print("✅ ЗАГРУЗКА УСПЕШНА!")
        else:
            print("❌ ЗАГРУЗКА НЕ УДАЛАСЬ")
        print("="*50)
        print(f"📊 Результат: {success_count}/{len(points)} точек")
        print(f"📚 Материал: {metadata.course_title}")
        print(f"🎓 Дисциплина: {metadata.discipline_name}")
        print(f"🆔 ID материала: {metadata.material_id}")
        
        return UploadResult(
            success=success_count > 0,
            points_uploaded=success_count,
            total_points=len(points),
            collection_name=self.collection_name,
            metadata=metadata,
            message=f"Загружено {success_count}/{len(points)} точек" if success_count > 0 else "Загрузка не удалась"
        )
    
    def _create_metadata_from_filename(self, file_path: str) -> DocumentMetadata:
        """Создает метаданные на основе имени файла"""
        metadata = DocumentMetadata()
        
        base_name = Path(file_path).stem
        base_name_display = base_name.replace('_', ' ').replace('-', ' ').title()
        
        metadata.discipline_name = base_name_display
        metadata.course_title = base_name_display
        
        metadata.discipline_id = self._get_simple_id(metadata.discipline_name)
        metadata.course_id = self._get_simple_id(metadata.course_title)
        metadata.material_id = self._get_simple_id(base_name)
        
        print(f"📋 Автоматически созданные метаданные:")
        print(f"   Дисциплина: {metadata.discipline_name}")
        print(f"   Курс: {metadata.course_title}")
        print(f"   ID материала: {metadata.material_id}")
        
        return metadata
    
    def upload_document_with_custom_metadata(self, doc_path: str, discipline_name: str, 
                                           course_title: str, difficulty: str = "medium") -> UploadResult:
        """Загрузка документа с кастомными метаданными"""
        metadata = DocumentMetadata(
            discipline_name=discipline_name,
            course_title=course_title,
            difficulty=difficulty
        )
        
        # Генерация ID
        metadata.discipline_id = self._get_simple_id(discipline_name)
        metadata.course_id = self._get_simple_id(course_title)
        metadata.material_id = self._get_simple_id(Path(doc_path).stem)
        
        return self.upload_document(doc_path, metadata)