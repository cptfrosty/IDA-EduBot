import uuid
import os
import re
import time
import hashlib
import sys
from pathlib import Path
from dataclasses import dataclass
from typing import List
import requests
import socket

import docx
from qdrant_client import QdrantClient, models
from qdrant_client.http.models import PointStruct
import numpy as np
from sentence_transformers import SentenceTransformer
import tqdm

# Конфигурация
QDRANT_HOST = "localhost"
QDRANT_PORT = 32770
COLLECTION_NAME = "IDA_edubot_materials"
VECTOR_SIZE = 1536
EMBEDDING_MODEL_NAME = "intfloat/multilingual-e5-large"

@dataclass
class DocumentMetadata:
    """Метаданные документа"""
    course_title: str = ""
    discipline_name: str = ""
    discipline_id: str = ""
    material_id: str = ""
    course_id: str = ""
    content_type: str = "document"
    difficulty: str = "medium"

def check_qdrant_connection(host: str, port: int, timeout: int = 5) -> bool:
    """Проверка подключения к Qdrant"""
    print(f"🔌 Проверка подключения к Qdrant {host}:{port}...")
    
    # Проверяем доступность порта
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(timeout)
        result = sock.connect_ex((host, port))
        sock.close()
        
        if result != 0:
            print(f"❌ Порт {port} недоступен на {host}")
            return False
    except Exception as e:
        print(f"❌ Ошибка подключения к порту: {e}")
        return False
    
    # Проверяем API Qdrant
    try:
        url = f"http://{host}:{port}"
        response = requests.get(f"{url}/", timeout=timeout)
        if response.status_code != 200:
            print(f"❌ Qdrant API не отвечает (статус {response.status_code})")
            return False
        
        print(f"✅ Qdrant доступен на {url}")
        return True
    except requests.exceptions.RequestException as e:
        print(f"❌ Не удалось подключиться к Qdrant API: {e}")
        return False
    except Exception as e:
        print(f"❌ Неизвестная ошибка: {e}")
        return False

class QdrantDocxUploader:
    def __init__(self, qdrant_host: str, qdrant_port: int):
        """Инициализация загрузчика"""
        print("🧠 Загрузка модели эмбеддингов...")
        try:
            self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
            print(f"✅ Модель загружена: {EMBEDDING_MODEL_NAME}")
            
            # Тестируем размерность
            test_embedding = self.embedding_model.encode(["тест"])[0]
            self.model_dim = len(test_embedding)
            print(f"📏 Размерность модели: {self.model_dim}")
            
            if self.model_dim != VECTOR_SIZE:
                print(f"⚠️  Размерность модели ({self.model_dim}) не совпадает с коллекцией ({VECTOR_SIZE})")
                print("   Будет выполнено дополнение/усечение векторов")
        except Exception as e:
            print(f"❌ Ошибка загрузки модели: {e}")
            sys.exit(1)
        
        # Подключение к Qdrant
        print(f"🔗 Подключение к Qdrant {qdrant_host}:{qdrant_port}...")
        try:
            self.client = QdrantClient(host=qdrant_host, port=qdrant_port, timeout=10)
            self.collection_name = COLLECTION_NAME
            
            # Проверяем коллекцию
            self._check_collection()
            
        except Exception as e:
            print(f"❌ Ошибка подключения к Qdrant: {e}")
            print("\nУбедитесь, что:")
            print(f"1. Qdrant запущен на {qdrant_host}:{qdrant_port}")
            print("2. Коллекция 'education_agent' существует")
            print("3. Порт открыт для подключения")
            sys.exit(1)
    
    def _check_collection(self):
        """Проверка существования и совместимости коллекции"""
        try:
            collections = self.client.get_collections().collections
            collection_names = [c.name for c in collections]
            
            if self.collection_name not in collection_names:
                print(f"❌ Коллекция '{self.collection_name}' не найдена")
                print("Доступные коллекции:", collection_names)
                sys.exit(1)
            
            collection_info = self.client.get_collection(self.collection_name)
            print(f"✅ Коллекция найдена: {self.collection_name}")
            print(f"📊 Текущее количество точек: {collection_info.points_count}")
            
            # Проверяем размерность векторов
            vectors_config = collection_info.config.params.vectors
            if hasattr(vectors_config, 'size'):
                collection_vector_size = vectors_config.size
            else:
                # Пытаемся получить размерность из структуры
                collection_vector_size = None
                if hasattr(vectors_config, 'values'):
                    for vec in vectors_config.values():
                        if hasattr(vec, 'size'):
                            collection_vector_size = vec.size
                            break
            
            if collection_vector_size:
                print(f"📐 Размерность коллекции: {collection_vector_size}")
                if collection_vector_size != self.model_dim:
                    print(f"⚠️  Внимание: модель ({self.model_dim}) ≠ коллекция ({collection_vector_size})")
            else:
                print("⚠️  Не удалось определить размерность коллекции")
            
        except Exception as e:
            print(f"❌ Ошибка проверки коллекции: {e}")
            sys.exit(1)
    
    def _pad_or_truncate_vector(self, vector: np.ndarray) -> np.ndarray:
        """Дополнение или усечение вектора до VECTOR_SIZE"""
        if len(vector) > VECTOR_SIZE:
            return vector[:VECTOR_SIZE]
        elif len(vector) < VECTOR_SIZE:
            padded = np.zeros(VECTOR_SIZE)
            padded[:len(vector)] = vector
            return padded
        else:
            return vector
    
    def _get_embeddings(self, texts: List[str]) -> List[np.ndarray]:
        """Получение эмбеддингов с приведением к нужной размерности"""
        embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
        processed_embeddings = [self._pad_or_truncate_vector(emb) for emb in embeddings]
        return processed_embeddings
    
    def _clean_text_preserve_structure(self, text: str) -> str:
        """Очистка текста с сохранением структуры"""
        # Убираем множественные пробелы
        text = re.sub(r'\s+', ' ', text)
        
        # Убираем странные символы, но сохраняем пунктуацию и кириллицу
        text = re.sub(r'[^\w\s.,!?;:()\-—\"\'\nа-яА-ЯёЁ]', '', text)
        
        # Восстанавливаем правильные пробелы после пунктуации
        text = re.sub(r'\s*([.,!?;:])\s*', r'\1 ', text)
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Разбивает текст на предложения"""
        sentences = []
        current_sentence = []
        i = 0
        
        while i < len(text):
            char = text[i]
            current_sentence.append(char)
            
            # Проверяем конец предложения
            if char in '.!?':
                # Проверяем, не является ли это точкой в сокращении
                if char == '.':
                    # Проверяем сокращения
                    abbreviations = ['т.д.', 'т.п.', 'др.', 'проф.', 'акад.', 'ст.', 'гл.',
                                    'рис.', 'табл.', 'см.', 'напр.', 'д.', 'и т.д.', 'и т.п.']
                    
                    is_abbreviation = False
                    for abbr in abbreviations:
                        if text[i-len(abbr)+1:i+1] == abbr:
                            is_abbreviation = True
                            break
                    
                    if not is_abbreviation and i + 1 < len(text) and text[i+1].isspace():
                        # Это конец предложения
                        sentence = ''.join(current_sentence).strip()
                        if sentence:
                            sentences.append(sentence)
                        current_sentence = []
                else:
                    # ! или ? - точно конец предложения
                    if i + 1 < len(text) and text[i+1].isspace():
                        sentence = ''.join(current_sentence).strip()
                        if sentence:
                            sentences.append(sentence)
                        current_sentence = []
            
            i += 1
        
        # Добавляем последнее предложение, если есть
        if current_sentence:
            sentence = ''.join(current_sentence).strip()
            if sentence:
                sentences.append(sentence)
        
        return sentences
    
    def _chunk_text_intelligent(self, text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
        """
        Умное разбиение текста на чанки с сохранением целостности предложений.
        """
        if not text or len(text.strip()) == 0:
            return []
        
        # Очистка текста
        text = self._clean_text_preserve_structure(text)
        
        # Если текст короче chunk_size, возвращаем его целиком
        if len(text) <= chunk_size:
            return [text]
        
        # Разбиваем на предложения
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Проверяем, начинается ли предложение с маленькой буквы (продолжение)
            starts_with_lower = sentence and sentence[0].islower()
            
            # Если текущий чанк пуст, а предложение начинается с маленькой буквы
            # это проблема - оно должно было быть частью предыдущего предложения
            if not current_chunk and starts_with_lower:
                # Начинаем чанк с этого предложения, но отмечаем проблему
                current_chunk = sentence
                continue
            
            # Проверяем, поместится ли предложение в текущий чанк
            if len(current_chunk) + len(sentence) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                # Текущий чанк заполнен, сохраняем его
                if current_chunk:
                    # Улучшаем конец чанка
                    current_chunk = self._improve_chunk_end(current_chunk)
                    chunks.append(current_chunk.strip())
                
                # Начинаем новый чанк
                current_chunk = sentence
        
        # Добавляем последний чанк
        if current_chunk:
            current_chunk = self._improve_chunk_end(current_chunk)
            chunks.append(current_chunk.strip())
        
        # Объединяем слишком маленькие чанки
        merged_chunks = []
        for chunk in chunks:
            if len(chunk) < 100 and merged_chunks:
                merged_chunks[-1] = merged_chunks[-1] + " " + chunk
            else:
                merged_chunks.append(chunk)
        
        # Добавляем перекрытие между чанками
        if len(merged_chunks) > 1:
            final_chunks = [merged_chunks[0]]
            
            for i in range(1, len(merged_chunks)):
                prev_chunk = merged_chunks[i-1]
                current_chunk = merged_chunks[i]
                
                # Берем конец предыдущего чанка для перекрытия
                if len(prev_chunk) > overlap:
                    overlap_text = prev_chunk[-overlap:]
                    # Находим начало последнего полного предложения в перекрытии
                    for delimiter in ['. ', '! ', '? ', '; ']:
                        pos = overlap_text.find(delimiter)
                        if pos != -1:
                            overlap_text = overlap_text[pos+2:]  # +2 чтобы пропустить разделитель и пробел
                            break
                else:
                    overlap_text = prev_chunk
                
                # Создаем новый чанк с перекрытием
                new_chunk = overlap_text + " " + current_chunk
                final_chunks.append(new_chunk.strip())
            
            return final_chunks
        
        return merged_chunks
    
    def _improve_chunk_end(self, chunk: str) -> str:
        """Улучшает конец чанка"""
        # Если чанк заканчивается на союз или предлог, пытаемся найти лучшее место
        bad_ends = [' и ', ' или ', ' но ', ' что ', ' который ', ' если ', 
                    ' когда ', ' где ', ' потому что ', ' поэтому ']
        
        for bad_end in bad_ends:
            bad_end_clean = bad_end.strip()
            if chunk.endswith(bad_end_clean):
                # Ищем последнюю хорошую точку разрыва
                for delimiter in ['.', '!', '?', ';', ',', ':', '-', '—']:
                    pos = chunk.rfind(delimiter)
                    if pos != -1 and pos > len(chunk) * 0.6:  # Не слишком близко к началу
                        return chunk[:pos+1]
                
                # Если не нашли, обрезаем плохой конец
                words = chunk.split()
                if len(words) > 2:
                    return ' '.join(words[:-1])
                break
        
        # Убедимся, что чанк заканчивается пунктуацией
        if chunk and chunk[-1] not in '.!?;':
            # Но не добавляем точку, если это список или число
            if not (chunk[-1].isdigit() or chunk.endswith(')') or 
                   chunk.endswith('"') or chunk.endswith("'")):
                chunk = chunk.rstrip(',:;-—') + '.'
        
        return chunk
    
    def _analyze_chunk_quality(self, chunks: List[str], original_text: str):
        """Анализ качества разбиения на чанки"""
        print("\n📊 АНАЛИЗ КАЧЕСТВА ЧАНКОВ:")
        print("-" * 50)
        
        total_chars = sum(len(c) for c in chunks)
        print(f"Всего чанков: {len(chunks)}")
        print(f"Общий размер: {total_chars:,} символов")
        print(f"Оригинальный текст: {len(original_text):,} символов")
        
        if len(original_text) > 0:
            loss = len(original_text) - total_chars
            print(f"Потеря данных: {loss:,} символов ({loss/len(original_text)*100:.1f}%)")
        
        # Проверяем начало чанков
        print("\n🔍 Проверка начала чанков:")
        bad_beginnings = 0
        for i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if chunk:
                # Проверяем, начинается ли чанк с маленькой буквы
                if chunk[0].islower():
                    # Но это может быть нормально для некоторых случаев (цитаты, продолжения)
                    bad_beginnings += 1
        
        # Проверяем концы чанков
        print("\n🔍 Проверка концов чанков:")
        bad_ends = 0
        for i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if chunk:
                # Проверяем пунктуацию в конце
                if chunk[-1] not in '.!?;':
                    # Но это нормально для списков, примеров и т.д.
                    if not (chunk.endswith(')') or chunk[-1].isdigit() or 
                           chunk.endswith('"') or chunk.endswith("'")):
                        bad_ends += 1
        
        print(f"\n📈 Статистика качества:")
        print(f"  Начинаются с маленькой буквы: {bad_beginnings} из {len(chunks)}")
        print(f"  Нет конечной пунктуации: {bad_ends} из {len(chunks)}")
        
        if len(chunks) > 0:
            quality_score = ((len(chunks) - (bad_beginnings + bad_ends) / 2) / len(chunks)) * 100
            print(f"  Общая оценка качества: {quality_score:.1f}%")
        
        # Показываем первые 2 чанка для проверки
        print("\n📝 Примеры чанков для проверки:")
        for i in range(min(2, len(chunks))):
            print(f"\n{'='*60}")
            print(f"ЧАНК {i+1} ({len(chunks[i]):,} символов):")
            print(f"{'='*60}")
            print(chunks[i][:500] + "..." if len(chunks[i]) > 500 else chunks[i])
            print(f"{'='*60}")
    
    def _read_docx(self, file_path: str) -> str:
        """Чтение docx файла"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Читаем таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"❌ Ошибка чтения docx: {e}")
            raise
    
    def _get_simple_id(self, text: str) -> str:
        """Простая генерация ID из строки"""
        hash_obj = hashlib.md5(text.lower().strip().encode())
        hex_dig = hash_obj.hexdigest()
        return f"{hex_dig[:8]}-{hex_dig[8:12]}-{hex_dig[12:16]}-{hex_dig[16:20]}-{hex_dig[20:]}"
    
    def _get_user_input(self, file_name: str) -> DocumentMetadata:
        """Запрос метаданных у пользователя"""
        print("\n" + "="*50)
        print("📋 МЕТАДАННЫЕ ДОКУМЕНТА")
        print("="*50)
        
        metadata = DocumentMetadata()
        
        # Используем имя файла как подсказку
        base_name = Path(file_name).stem
        base_name = base_name.replace('_', ' ').replace('-', ' ').title()
        
        print(f"📄 Имя файла: {base_name}")
        
        # Дисциплина
        default_discipline = input(f"Название дисциплины [{base_name}]: ").strip()
        if not default_discipline:
            default_discipline = base_name
        metadata.discipline_name = default_discipline
        
        # Курс
        default_course = input(f"Название курса [{default_discipline}]: ").strip()
        metadata.course_title = default_course if default_course else default_discipline
        
        # Материал
        material_name = input(f"Название материала [{base_name}]: ").strip() or base_name
        
        # Генерация ID
        print("\n🔧 Генерация ID...")
        metadata.discipline_id = self._get_simple_id(metadata.discipline_name)
        metadata.course_id = self._get_simple_id(metadata.course_title)
        metadata.material_id = self._get_simple_id(material_name)
        
        # Уровень сложности
        print("\n📊 Уровень сложности:")
        print("1. beginner - начальный")
        print("2. medium - средний (по умолчанию)")
        print("3. advanced - продвинутый")
        
        diff_map = {"1": "beginner", "2": "medium", "3": "advanced"}
        diff_choice = input("Выберите (1-3) [2]: ").strip() or "2"
        metadata.difficulty = diff_map.get(diff_choice, "medium")
        
        metadata.content_type = "document"
        
        print("\n" + "="*50)
        print("✅ Метаданные сохранены")
        print("="*50)
        
        return metadata
    
    def upload_document(self, doc_path: str):
        """Загрузка документа в коллекцию"""
        if not os.path.exists(doc_path):
            print(f"❌ Файл не найден: {doc_path}")
            return
        
        print(f"\n📄 Обработка документа: {Path(doc_path).name}")
        
        # Получение метаданных
        metadata = self._get_user_input(doc_path)
        
        # Чтение документа
        print("📖 Чтение документа...")
        try:
            full_text = self._read_docx(doc_path)
            print(f"   Прочитано символов: {len(full_text):,}")
            
            if len(full_text.strip()) < 50:
                print("❌ Текст слишком короткий")
                return
        except Exception as e:
            print(f"❌ Ошибка чтения файла: {e}")
            return
        
        # Разбивка на чанки
        print("✂️  Интеллектуальное разбиение на чанки...")
        chunks = self._chunk_text_intelligent(full_text, chunk_size=800, overlap=200)
        print(f"   Создано чанков: {len(chunks)}")
        
        if not chunks:
            print("❌ Не удалось создать чанки")
            return
        
        # Анализ качества чанков
        self._analyze_chunk_quality(chunks, full_text)
        
        # Спрашиваем подтверждение
        print("\n⚠️  ПРЕДУПРЕЖДЕНИЕ: Проверьте качество чанков выше.")
        choice = input("Продолжить загрузку? (y/n): ").strip().lower()
        if choice != 'y':
            print("Загрузка отменена пользователем")
            return
        
        # Генерация эмбеддингов
        print("🔢 Генерация эмбеддингов...")
        try:
            embeddings = self._get_embeddings(chunks)
            print(f"   Сгенерировано эмбеддингов: {len(embeddings)}")
        except Exception as e:
            print(f"❌ Ошибка генерации эмбеддингов: {e}")
            return
        
        # Подготовка точек
        print("📦 Подготовка данных...")
        points = []
        
        # Получаем текущее количество точек для продолжения нумерации
        try:
            existing_count = self.client.count(
                collection_name=self.collection_name,
                exact=True
            ).count
            start_id = existing_count + 1
        except:
            start_id = 1
        
        for idx, (chunk_text, embedding) in enumerate(tqdm.tqdm(zip(chunks, embeddings), 
                                                                total=len(chunks), 
                                                                desc="Подготовка")):
            # ID точки
            point_id = start_id + idx
            
            # ID чанка
            chunk_hash = hashlib.md5(f"{metadata.material_id}_{idx}".encode()).hexdigest()[:8]
            chunk_id_int = int(chunk_hash, 16) % 1000000
            
            point = PointStruct(
                id=point_id,
                vector=embedding.tolist(),
                payload={
                    "content_type": metadata.content_type,
                    "discipline_id": metadata.discipline_id,
                    "chunk_id": chunk_id_int,
                    "chunk_index": idx,
                    "course_title": metadata.course_title,
                    "difficulty": metadata.difficulty,
                    "material_id": metadata.material_id,
                    "course_id": metadata.course_id,
                    "discipline_name": metadata.discipline_name,
                    "chunk_text": chunk_text[:5000],
                    "source_file": os.path.basename(doc_path),
                    "total_chunks": len(chunks),
                    "upload_timestamp": int(time.time())
                }
            )
            points.append(point)
        
        # Загрузка в Qdrant
        print(f"\n🚀 Загрузка {len(points)} точек в Qdrant...")
        
        batch_size = 10
        success_count = 0
        
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            try:
                self.client.upsert(
                    collection_name=self.collection_name,
                    points=batch,
                    wait=True
                )
                success_count += len(batch)
                print(f"   ✓ Батч {i//batch_size + 1}: загружено {min(i + batch_size, len(points))}/{len(points)}")
            except Exception as e:
                print(f"   ✗ Ошибка батча {i//batch_size + 1}: {str(e)[:200]}")
                # Пробуем загрузить по одной
                for j, point in enumerate(batch):
                    try:
                        self.client.upsert(
                            collection_name=self.collection_name,
                            points=[point],
                            wait=False
                        )
                        success_count += 1
                    except Exception as point_error:
                        print(f"      ✗ Ошибка точки {i + j}: {str(point_error)[:100]}")
        
        # Итоги
        print("\n" + "="*50)
        if success_count > 0:
            print("✅ ЗАГРУЗКА УСПЕШНА!")
        else:
            print("❌ ЗАГРУЗКА НЕ УДАЛАСЬ")
        print("="*50)
        print(f"📊 Результат: {success_count}/{len(points)} точек")
        print(f"📚 Материал: {metadata.course_title}")
        print(f"🎓 Дисциплина: {metadata.discipline_name}")
        print(f"🆔 ID материала: {metadata.material_id}")
        print(f"📈 Уровень: {metadata.difficulty}")
        
        # Обновленная информация о коллекции
        self.show_collection_stats()
    
    def show_collection_stats(self):
        """Показать статистику коллекции"""
        try:
            info = self.client.get_collection(self.collection_name)
            print(f"\n📊 Коллекция '{self.collection_name}':")
            print(f"   Всего точек: {info.points_count:,}")
            print(f"   Проиндексировано векторов: {info.indexed_vectors_count:,}")
            print(f"   Статус: {info.status}")
            
            if hasattr(info, 'segments_count'):
                print(f"   Сегментов: {info.segments_count}")
            
        except Exception as e:
            print(f"⚠️  Не удалось получить статистику: {e}")
    
    def test_search(self):
        """Тестовый поиск для проверки работы"""
        try:
            print("\n🔍 Тестовый поиск...")
            results = self.client.scroll(
                collection_name=self.collection_name,
                limit=2,
                with_payload=True,
                with_vectors=False
            )
            
            if results[0]:
                print("✅ Данные в коллекции найдены:")
                for point in results[0]:
                    print(f"  ID: {point.id}")
                    print(f"  Дисциплина: {point.payload.get('discipline_name', 'N/A')}")
                    print(f"  Чанк: {point.payload.get('chunk_index', 'N/A')}")
                    print(f"  Текст: {point.payload.get('chunk_text', '')[:80]}...")
                    print()
            else:
                print("ℹ️  Коллекция пуста")
        except Exception as e:
            print(f"❌ Ошибка поиска: {e}")

def main():
    """Основная функция"""
    print("="*60)
    print("📚 ЗАГРУЗЧИК DOCX ДОКУМЕНТОВ В QDRANT")
    print("="*60)
    
    # Проверка зависимостей
    try:
        import docx
        import numpy as np
        from sentence_transformers import SentenceTransformer
        from qdrant_client import QdrantClient
    except ImportError as e:
        print(f"❌ Не установлены зависимости: {e}")
        print("\n📦 Установите зависимости:")
        print("pip install python-docx sentence-transformers numpy qdrant-client tqdm")
        sys.exit(1)
    
    # Настройки
    qdrant_host = QDRANT_HOST
    qdrant_port = QDRANT_PORT
    
    print(f"📍 Целевой хост: {qdrant_host}:{qdrant_port}")
    
    # Проверка подключения к Qdrant
    if not check_qdrant_connection(qdrant_host, qdrant_port):
        print("\n❌ Не удалось подключиться к Qdrant.")
        print("Пожалуйста, выполните следующие шаги:")
        print("1. Убедитесь, что Qdrant запущен")
        print("2. Проверьте порт подключения")
        print("3. Убедитесь, что нет фаервола, блокирующего подключение")
        print("\nЗапуск Qdrant через Docker:")
        print("docker run -p 6333:6333 qdrant/qdrant")
        sys.exit(1)
    
    # Инициализация загрузчика
    print("\n🚀 Инициализация загрузчика...")
    try:
        uploader = QdrantDocxUploader(qdrant_host, qdrant_port)
        print("✅ Загрузчик инициализирован")
    except Exception as e:
        print(f"❌ Ошибка инициализации: {e}")
        return
    
    # Основной цикл
    while True:
        print("\n" + "-"*50)
        print("🏠 ГЛАВНОЕ МЕНЮ")
        print("-"*50)
        print("1. 📤 Загрузить документ")
        print("2. 📊 Показать статистику коллекции")
        print("3. 🔍 Проверить данные в коллекции")
        print("4. 🧪 Тестовый поиск")
        print("5. 🚪 Выход")
        
        choice = input("\n🎯 Выберите действие (1-5): ").strip()
        
        if choice == "1":
            doc_path = input("📁 Введите путь к .docx файлу: ").strip()
            if doc_path:
                if not doc_path.endswith('.docx'):
                    print("⚠️  Файл должен иметь расширение .docx")
                    continue
                if not os.path.exists(doc_path):
                    print("❌ Файл не найден")
                    continue
                uploader.upload_document(doc_path)
            else:
                print("❌ Путь не может быть пустым")
        
        elif choice == "2":
            uploader.show_collection_stats()
        
        elif choice == "3":
            try:
                result = uploader.client.count(
                    collection_name=uploader.collection_name,
                    exact=True
                )
                print(f"\n📈 Всего точек в коллекции: {result.count}")
            except Exception as e:
                print(f"❌ Ошибка: {e}")
        
        elif choice == "4":
            uploader.test_search()
        
        elif choice == "5":
            print("\n👋 Выход из программы")
            break
        
        else:
            print("❌ Неверный выбор. Попробуйте снова.")

if __name__ == "__main__":
    main()