# file: backend/app/qdrant_integration/qdrant_loader.py

import uuid
import os
import re
import time
import hashlib
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Optional, Dict, Any
import logging

import docx
from qdrant_client import QdrantClient, models
from qdrant_client.http.models import PointStruct, VectorParams, Distance
import numpy as np
from sentence_transformers import SentenceTransformer
import tqdm

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Метаданные документа"""
    course_title: str = ""
    discipline_name: str = ""
    discipline_id: str = ""
    material_id: str = ""
    course_id: str = ""
    content_type: str = "document"
    difficulty: str = "medium"
    
    def to_dict(self) -> Dict[str, Any]:
        """Конвертация в словарь"""
        return asdict(self)

class QdrantDocumentUploader:
    """Класс для загрузки документов в Qdrant"""
    
    def __init__(
        self, 
        qdrant_host: str = "localhost", 
        qdrant_port: int = 6333,
        collection_name: str = "ida_edubot",
        embedding_model_name: str = "intfloat/multilingual-e5-large"
    ):
        """
        Инициализация загрузчика документов
        
        Args:
            qdrant_host: Хост Qdrant
            qdrant_port: Порт Qdrant
            collection_name: Имя коллекции
            embedding_model_name: Название модели эмбеддингов
        """
        self.qdrant_host = qdrant_host
        self.qdrant_port = qdrant_port
        self.collection_name = collection_name
        self.embedding_model_name = embedding_model_name
        
        # Инициализация модели и клиента
        self._initialize_embedding_model()
        self._initialize_qdrant_client()
        
    def _initialize_embedding_model(self):
        """Инициализация модели эмбеддингов"""
        try:
            logger.info(f"Загрузка модели эмбеддингов: {self.embedding_model_name}")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            
            # Тестируем модель
            test_embedding = self.embedding_model.encode(["тест"])[0]
            self.model_dim = len(test_embedding)
            
            logger.info(f"Модель загружена. Размерность: {self.model_dim}")
            
        except Exception as e:
            logger.error(f"Ошибка загрузки модели эмбеддингов: {e}")
            raise
    
    def _initialize_qdrant_client(self):
        """Инициализация клиента Qdrant"""
        try:
            logger.info(f"Подключение к Qdrant {self.qdrant_host}:{self.qdrant_port}")
            self.client = QdrantClient(
                host=self.qdrant_host, 
                port=self.qdrant_port, 
                timeout=30
            )
            
            # Проверяем соединение
            self.client.get_collections()
            logger.info("Подключение к Qdrant успешно")
            
        except Exception as e:
            logger.error(f"Ошибка подключения к Qdrant: {e}")
            raise
    
    def ensure_collection_exists(self, vector_size: Optional[int] = None) -> bool:
        """
        Убедиться, что коллекция существует
        
        Args:
            vector_size: Размерность векторов (если None, используется model_dim)
            
        Returns:
            True если коллекция существует или создана
        """
        try:
            collections = self.client.get_collections().collections
            collection_names = [c.name for c in collections]
            
            if self.collection_name in collection_names:
                logger.info(f"Коллекция '{self.collection_name}' уже существует")
                return True
            
            # Создаем коллекцию если ее нет
            if vector_size is None:
                vector_size = self.model_dim
            
            logger.info(f"Создание коллекции '{self.collection_name}' с размерностью {vector_size}")
            
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(
                    size=vector_size,
                    distance=Distance.COSINE
                )
            )
            
            # Создаем индексы
            self._create_payload_indexes()
            
            logger.info(f"Коллекция '{self.collection_name}' создана")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка при работе с коллекцией: {e}")
            return False
    
    def _create_payload_indexes(self):
        """Создание индексов для полей payload"""
        index_fields = [
            ("discipline_id", models.PayloadSchemaType.KEYWORD),
            ("course_id", models.PayloadSchemaType.KEYWORD),
            ("material_id", models.PayloadSchemaType.KEYWORD),
            ("discipline_name", models.PayloadSchemaType.TEXT),
            ("course_title", models.PayloadSchemaType.TEXT),
            ("difficulty", models.PayloadSchemaType.KEYWORD),
            ("content_type", models.PayloadSchemaType.KEYWORD),
        ]
        
        for field_name, schema_type in index_fields:
            try:
                self.client.create_payload_index(
                    collection_name=self.collection_name,
                    field_name=field_name,
                    field_schema=schema_type,
                )
                logger.debug(f"Создан индекс для поля '{field_name}'")
            except Exception as e:
                if "already exists" not in str(e):
                    logger.warning(f"Не удалось создать индекс для '{field_name}': {e}")
    
    def _read_docx(self, file_path: str) -> str:
        """Чтение DOCX файла"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Читаем параграфы
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Читаем таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
            
        except Exception as e:
            logger.error(f"Ошибка чтения DOCX файла: {e}")
            raise
    
    def _chunk_text(self, text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
        """
        Разбивка текста на чанки
        
        Args:
            text: Исходный текст
            chunk_size: Размер чанка в символах
            overlap: Перекрытие между чанками
            
        Returns:
            Список чанков
        """
        if not text:
            return []
        
        # Очистка текста
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Простое разбиение на предложения
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Начинаем новый чанк с частью предыдущего для перекрытия
                if overlap > 0 and chunks:
                    last_chunk = chunks[-1]
                    overlap_text = last_chunk[-overlap:] if len(last_chunk) > overlap else last_chunk
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence
        
        # Добавляем последний чанк
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _get_embeddings(self, texts: List[str]) -> List[np.ndarray]:
        """Получение эмбеддингов для текстов"""
        try:
            embeddings = self.embedding_model.encode(
                texts, 
                show_progress_bar=False,
                normalize_embeddings=True
            )
            return embeddings
        except Exception as e:
            logger.error(f"Ошибка генерации эмбеддингов: {e}")
            raise
    
    def upload_document(
        self, 
        file_path: str, 
        metadata: DocumentMetadata,
        chunk_size: int = 800,
        overlap: int = 200
    ) -> Dict[str, Any]:
        """
        Загрузка документа в Qdrant
        
        Args:
            file_path: Путь к файлу
            metadata: Метаданные документа
            chunk_size: Размер чанка
            overlap: Перекрытие между чанками
            
        Returns:
            Словарь с результатом загрузки
        """
        start_time = time.time()
        
        try:
            # Проверяем файл
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Файл не найден: {file_path}")
            
            # Убеждаемся, что коллекция существует
            if not self.ensure_collection_exists():
                raise Exception("Не удалось создать или найти коллекцию")
            
            # Читаем документ
            logger.info(f"Чтение документа: {file_path}")
            text = self._read_docx(file_path)
            
            if len(text.strip()) < 50:
                raise ValueError("Текст документа слишком короткий")
            
            logger.info(f"Прочитано {len(text)} символов")
            
            # Разбиваем на чанки
            logger.info(f"Разбиение на чанки (размер: {chunk_size}, перекрытие: {overlap})")
            chunks = self._chunk_text(text, chunk_size, overlap)
            
            if not chunks:
                raise ValueError("Не удалось создать чанки из текста")
            
            logger.info(f"Создано {len(chunks)} чанков")
            
            # Генерируем эмбеддинги
            logger.info("Генерация эмбеддингов...")
            embeddings = self._get_embeddings(chunks)
            
            # Подготавливаем точки для загрузки
            points = []
            existing_count = self.client.count(
                collection_name=self.collection_name,
                exact=True
            ).count
            
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                point_id = existing_count + idx + 1
                
                # Генерируем уникальный ID чанка
                chunk_hash = hashlib.md5(
                    f"{metadata.material_id}_{idx}_{chunk_text[:100]}".encode()
                ).hexdigest()[:16]
                
                point = PointStruct(
                    id=point_id,
                    vector=embedding.tolist(),
                    payload={
                        "content_type": metadata.content_type,
                        "discipline_id": metadata.discipline_id,
                        "course_title": metadata.course_title,
                        "difficulty": metadata.difficulty,
                        "material_id": metadata.material_id,
                        "course_id": metadata.course_id,
                        "discipline_name": metadata.discipline_name,
                        "chunk_text": chunk_text,
                        "chunk_index": idx,
                        "chunk_id": chunk_hash,
                        "total_chunks": len(chunks),
                        "source_file": os.path.basename(file_path),
                        "upload_timestamp": int(time.time())
                    }
                )
                points.append(point)
            
            # Загружаем в Qdrant батчами
            logger.info(f"Загрузка {len(points)} точек в Qdrant...")
            
            batch_size = 50
            success_count = 0
            
            for i in range(0, len(points), batch_size):
                batch = points[i:i + batch_size]
                try:
                    self.client.upsert(
                        collection_name=self.collection_name,
                        points=batch,
                        wait=True
                    )
                    success_count += len(batch)
                    logger.debug(f"Загружен батч {i//batch_size + 1}")
                except Exception as e:
                    logger.error(f"Ошибка при загрузке батча {i//batch_size + 1}: {e}")
                    # Пробуем загрузить по одной точке
                    for point in batch:
                        try:
                            self.client.upsert(
                                collection_name=self.collection_name,
                                points=[point],
                                wait=False
                            )
                            success_count += 1
                        except:
                            pass
            
            # Получаем итоговую статистику
            final_count = self.client.count(
                collection_name=self.collection_name,
                exact=True
            ).count
            
            elapsed_time = time.time() - start_time
            
            result = {
                "success": True,
                "file_name": os.path.basename(file_path),
                "material_id": metadata.material_id,
                "course_title": metadata.course_title,
                "discipline_name": metadata.discipline_name,
                "stats": {
                    "total_chunks": len(chunks),
                    "points_uploaded": success_count,
                    "total_points_in_collection": final_count,
                    "processing_time_seconds": round(elapsed_time, 2)
                },
                "metadata": metadata.to_dict()
            }
            
            logger.info(f"Документ успешно загружен. Загружено {success_count} точек")
            return result
            
        except Exception as e:
            elapsed_time = time.time() - start_time
            logger.error(f"Ошибка загрузки документа: {e}")
            
            return {
                "success": False,
                "error": str(e),
                "processing_time_seconds": round(elapsed_time, 2),
                "file_path": file_path
            }
    
    def search_documents(
        self,
        query: str,
        limit: int = 5,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Поиск документов по запросу
        
        Args:
            query: Поисковый запрос
            limit: Количество результатов
            filters: Фильтры поиска
            
        Returns:
            Список найденных документов
        """
        try:
            # Генерируем эмбеддинг для запроса
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Подготавливаем фильтры
            search_filters = None
            if filters:
                filter_conditions = []
                for key, value in filters.items():
                    if value:
                        filter_conditions.append(
                            models.FieldCondition(
                                key=key,
                                match=models.MatchValue(value=value)
                            )
                        )
                
                if filter_conditions:
                    search_filters = models.Filter(
                        must=filter_conditions
                    )
            
            # Выполняем поиск
            search_result = self.client.search(
                collection_name=self.collection_name,
                query_vector=query_embedding.tolist(),
                limit=limit,
                query_filter=search_filters,
                with_payload=True,
                with_vectors=False
            )
            
            # Форматируем результаты
            results = []
            for hit in search_result:
                results.append({
                    "id": hit.id,
                    "score": hit.score,
                    "payload": hit.payload,
                    "text_preview": hit.payload.get("chunk_text", "")[:200] + "..."
                })
            
            return results
            
        except Exception as e:
            logger.error(f"Ошибка поиска: {e}")
            return []
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Получение статистики коллекции"""
        try:
            collection_info = self.client.get_collection(self.collection_name)
            
            return {
                "name": self.collection_name,
                "points_count": collection_info.points_count,
                "indexed_vectors_count": collection_info.indexed_vectors_count,
                "status": collection_info.status
            }
        except Exception as e:
            logger.error(f"Ошибка получения статистики: {e}")
            return {}